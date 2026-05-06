"""Streamlit front‑end for the legal monitor."""

from __future__ import annotations

import datetime as dt
import io
import os
import re
from typing import Any, Dict, List

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

import monitor

st.set_page_config(page_title="Legal Monitor", layout="wide")

st.title("Liquidation & Bankruptcy Monitor")

if "last_summary" not in st.session_state:
    st.session_state["last_summary"] = None
if "last_records" not in st.session_state:
    st.session_state["last_records"] = []


def _flatten_record(record: Dict[str, Any]) -> Dict[str, Any]:
    """Flatten nested JSON into table-friendly keys."""
    flattened: Dict[str, Any] = {}

    def walk(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for k, v in value.items():
                walk(f"{prefix}.{k}" if prefix else str(k), v)
        elif isinstance(value, list):
            flattened[prefix] = ", ".join(str(item) for item in value)
        else:
            flattened[prefix] = value

    walk("", record)
    return flattened


def _records_to_dataframe(records: List[Dict[str, Any]]) -> pd.DataFrame:
    if not records:
        return pd.DataFrame()
    return pd.DataFrame([_flatten_record(rec) for rec in records]).fillna("")


def _export_excel(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Results")
    return output.getvalue()


def _record_sections(record: Dict[str, Any]) -> List[tuple[str, Any]]:
    debtor = record.get("debtor") if isinstance(record.get("debtor"), dict) else {}
    proposer_names: List[str] = []
    proposers = record.get("proposers")
    if isinstance(proposers, list):
        for proposer in proposers:
            if isinstance(proposer, dict):
                name = proposer.get("corporate_body_name")
                if name:
                    proposer_names.append(str(name))
            elif proposer:
                proposer_names.append(str(proposer))

    return [
        ("ID", record.get("id", "")),
        ("Dataset", record.get("dataset", "")),
        ("Court", record.get("court_name", "")),
        ("File reference", record.get("file_reference", "")),
        ("Kind", record.get("kind", "")),
        ("Released date", record.get("released_date", "")),
        ("Updated at", record.get("updated_at", "")),
        ("Debtor company", debtor.get("corporate_body_name") or record.get("corporate_body_name", "")),
        ("Company CIN", debtor.get("cin") or record.get("cin", "")),
        ("Proposers", "; ".join(proposer_names)),
        ("Heading", record.get("heading", "")),
        ("Decision", record.get("decision", "")),
        ("Announcement", record.get("announcement", "")),
        ("Advice", record.get("advice", "")),
    ]


def _format_date(value: Any) -> str:
    if not value:
        return "-"
    text = str(value)
    try:
        parsed = dt.datetime.fromisoformat(text.replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            return parsed.strftime("%d %B %Y")
        return parsed.astimezone(dt.timezone.utc).strftime("%d %B %Y, %H:%M UTC")
    except ValueError:
        return text


def _build_template_placeholders(records: List[Dict[str, Any]], summary: Dict[str, Any]) -> Dict[str, str]:
    first = records[0] if records else {}
    debtor = first.get("debtor") if isinstance(first.get("debtor"), dict) else {}
    entity_name = debtor.get("corporate_body_name") or first.get("corporate_body_name", "-")
    proposer = ""
    proposers = first.get("proposers")
    if isinstance(proposers, list) and proposers:
        candidate = proposers[0]
        if isinstance(candidate, dict):
            proposer = str(candidate.get("corporate_body_name", ""))
        else:
            proposer = str(candidate)

    finding_headline = first.get("heading") or first.get("kind") or "No finding"
    notice_type = first.get("kind") or first.get("dataset") or "-"
    source_extracts = []
    for idx, rec in enumerate(records[:10], 1):
        excerpt = rec.get("announcement") or rec.get("decision") or rec.get("advice") or rec.get("heading") or "-"
        source_extracts.append(f"{idx}. {excerpt}")

    management_summary = (
        f"Monitoring for '{summary.get('query', '-')}' found {summary.get('matches', 0)} relevant notice(s) "
        f"from {summary.get('fetched', 0)} reviewed records in the selected period."
    )
    if records:
        management_summary += f" The latest highlighted record is '{finding_headline}' for entity '{entity_name}'."

    return {
        "{{GENERATED_AT}}": dt.datetime.utcnow().strftime("%d %B %Y, %H:%M UTC"),
        "{{SEARCH_TARGET}}": str(summary.get("query", "-")),
        "{{PERIOD_FROM}}": _format_date(summary.get("since")),
        "{{PERIOD_TO}}": _format_date(summary.get("to")),
        "{{RECORDS_REVIEWED}}": str(summary.get("fetched", 0)),
        "{{RELEVANT_FINDINGS}}": str(summary.get("matches", 0)),
        "{{FINDING_HEADLINE}}": str(finding_headline),
        "{{ENTITY_NAME}}": str(entity_name),
        "{{NOTICE_TYPE}}": str(notice_type),
        "{{COURT_NAME}}": str(first.get("court_name", "-")),
        "{{REFERENCE}}": str(first.get("file_reference", "-")),
        "{{PUBLICATION_DATE}}": _format_date(first.get("released_date")),
        "{{UPDATED_AT}}": _format_date(first.get("updated_at")),
        "{{RELEVANT_PARTY}}": proposer or str(entity_name),
        "{{MANAGEMENT_SUMMARY}}": management_summary,
        "{{SOURCE_EXTRACT}}": "\n".join(source_extracts) if source_extracts else "No findings in selected period.",
    }


def _replace_placeholders_in_doc(doc: Document, replacements: Dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, val)


def _export_word(df: pd.DataFrame, records: List[Dict[str, Any]], summary: Dict[str, Any]) -> bytes:
    template_path = os.path.join("docs", "legal_monitor_template.docx")
    doc = Document(template_path) if os.path.exists(template_path) else Document()
    if not os.path.exists(template_path):
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10)
    _replace_placeholders_in_doc(doc, _build_template_placeholders(records, summary))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _export_pdf(df: pd.DataFrame, records: List[Dict[str, Any]], summary: Dict[str, Any]) -> bytes:
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    width, height = A4

    def draw_wrapped(text: str, x: int, y: float, font_name: str = "Helvetica", font_size: int = 9) -> float:
        max_width = width - x - 40
        lines = simpleSplit(text, font_name, font_size, max_width)
        c.setFont(font_name, font_size)
        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 40
                c.setFont(font_name, font_size)
            c.drawString(x, y, line)
            y -= 12
        return y

    repl = _build_template_placeholders(records, summary)
    y = height - 40
    c.setFont("Helvetica-Bold", 13)
    c.drawString(40, y, "Legal Monitor Briefing")
    y = draw_wrapped("Public registry monitoring summary for senior management", 40, y - 20)
    y -= 8
    c.setFont("Helvetica-Bold", 10)
    y = draw_wrapped("Executive overview", 40, y)
    c.setFont("Helvetica", 9)
    overview_lines = [
        f"Monitored entity: {repl['{{SEARCH_TARGET}}']}",
        f"Reporting period: {repl['{{PERIOD_FROM}}']} - {repl['{{PERIOD_TO}}']}",
        f"Report generated: {repl['{{GENERATED_AT}}']}",
        f"Records reviewed: {repl['{{RECORDS_REVIEWED}}']}",
        f"Relevant findings: {repl['{{RELEVANT_FINDINGS}}']}",
        f"Key point: {repl['{{MANAGEMENT_SUMMARY}}']}",
    ]
    for line in overview_lines:
        y = draw_wrapped(line, 40, y - 2)
    y -= 6
    c.setFont("Helvetica-Bold", 10)
    y = draw_wrapped("Finding details", 40, y)
    details = [
        ("Headline", repl["{{FINDING_HEADLINE}}"]),
        ("Entity", repl["{{ENTITY_NAME}}"]),
        ("Notice type", repl["{{NOTICE_TYPE}}"]),
        ("Court", repl["{{COURT_NAME}}"]),
        ("Reference", repl["{{REFERENCE}}"]),
        ("Publication date", repl["{{PUBLICATION_DATE}}"]),
        ("Latest registry update", repl["{{UPDATED_AT}}"]),
        ("Relevant party", repl["{{RELEVANT_PARTY}}"]),
    ]
    c.setFont("Helvetica", 9)
    for label, value in details:
        y = draw_wrapped(f"{label}: {value}", 40, y - 2)
    y -= 4
    c.setFont("Helvetica-Bold", 10)
    y = draw_wrapped("Source extract", 40, y)
    c.setFont("Helvetica", 8)
    y = draw_wrapped(repl["{{SOURCE_EXTRACT}}"], 40, y - 2)

    c.save()
    return output.getvalue()


def _build_export_basename(search_query: str) -> str:
    cleaned_query = re.sub(r"[^A-Za-z0-9]", "", (search_query or ""))
    query_part = (cleaned_query[:10] if cleaned_query else "NoQuery")
    timestamp = dt.datetime.utcnow()
    return f"LegalMonitor_{query_part}_{timestamp:%d%m%Y}_{timestamp:%H%M%S}"


search_query = st.text_input(
    "What name or keyword should we look for?",
    value="",
    help=(
        "Type a company name, person name, or keyword. "
        "We look for this text without case sensitivity "
        "(for example, 'Acme' matches 'ACME')."
    ),
)
search_scope = st.radio(
    "Where should we look for that text?",
    options=[
        ("full_text", "Every available field"),
        ("targeted", "Only names (proposers and corporate body name)"),
        ("combined", "Names first, then all other fields"),
    ],
    format_func=lambda item: item[1],
    horizontal=False,
    help=(
        "This setting controls which data fields are checked in each API record. "
        "'Names first, then all other fields' starts with name fields and then checks the rest."
    ),
)
window_mode = st.radio(
    "How far back should we search?",
    options=["Last N days", "Custom date range"],
    horizontal=True,
    help=(
        "Choose a rolling period (for example, last 30 days) "
        "or set exact start and end dates in UTC."
    ),
)

if window_mode == "Last N days":
    trailing_days = st.number_input(
        "Number of days to include",
        min_value=0,
        value=30,
        step=1,
        help="Example: 30 means from today back to 30 days ago (UTC).",
    )
    date_from = None
    date_to = None
else:
    trailing_days = None
    default_from = (dt.datetime.utcnow() - dt.timedelta(days=30)).date()
    date_from = st.date_input(
        "Start date (UTC)",
        value=default_from,
        help="Beginning of the period to check.",
    )
    date_to = st.date_input(
        "End date (UTC)",
        value=dt.datetime.utcnow().date(),
        help="Final date to include in the search period.",
    )

current_dir = os.path.dirname(os.path.abspath(__file__))
ts_path = os.path.join(current_dir, "last_run.txt")

last_ts = monitor.load_last_run_timestamp(ts_path)
if last_ts:
    last_dt = dt.datetime.fromisoformat(last_ts.rstrip("Z"))
    st.write(f"Last update run: {last_dt.strftime('%Y-%m-%d %H:%M:%S')} UTC")
else:
    st.write("No previous runs recorded.")

if st.button("Run update"):
    with st.spinner("Fetching updates..."):
        if window_mode == "Last N days":
            summary = monitor.perform_update_last_n_days(
                int(trailing_days),
                query=search_query,
                search_mode=search_scope[0],
                send_notifications=False,
            )
        else:
            since_dt = dt.datetime.combine(date_from, dt.time.min)
            to_dt = dt.datetime.combine(date_to, dt.time.max)
            since = since_dt.isoformat() + "Z"
            to_timestamp = to_dt.isoformat() + "Z"
            summary = monitor.perform_update(
                since,
                query=search_query,
                search_mode=search_scope[0],
                send_notifications=False,
                to_timestamp=to_timestamp,
            )

        monitor.save_last_run_timestamp(ts_path, summary["timestamp"])

    st.session_state["last_summary"] = summary
    st.session_state["last_records"] = summary.get("records", [])
    st.success("Update complete.")

if st.button("Clear results"):
    st.session_state["last_summary"] = None
    st.session_state["last_records"] = []
    st.success("Cleared previous search results.")

summary = st.session_state.get("last_summary")
records = st.session_state.get("last_records", [])
if summary is not None:
    df = _records_to_dataframe(records)
    export_basename = _build_export_basename(summary.get("query", ""))

    st.subheader("Run Summary")
    st.write({k: v for k, v in summary.items() if k != "records"})

    st.subheader("Matching Results")
    if df.empty:
        st.info("No matching results found for the selected filters.")
    else:
        st.dataframe(df, use_container_width=True)

    st.subheader("Download Results")
    st.download_button(
        "Download XLSX",
        data=_export_excel(df),
        file_name=f"{export_basename}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    st.download_button(
        "Download Word",
        data=_export_word(df, records, summary),
        file_name=f"{export_basename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.download_button(
        "Download PDF",
        data=_export_pdf(df, records, summary),
        file_name=f"{export_basename}.pdf",
        mime="application/pdf",
    )

    st.subheader("Email Results")
    recipient_input = st.text_input(
        "Recipients (comma-separated emails)",
        value="",
        help="No email is sent automatically. Click the button below to send.",
    )
    recipients = [item.strip() for item in recipient_input.split(",") if item.strip()]
    attachment_options = st.multiselect(
        "Attachments",
        options=["xlsx", "pdf", "word"],
        default=["xlsx"],
        help="Choose one or more file formats to attach to the email.",
    )
    if st.button("Send via email"):
        attachments_payload: List[Dict[str, str | bytes]] = []
        if "xlsx" in attachment_options:
            attachments_payload.append(
                {
                    "filename": f"{export_basename}.xlsx",
                    "content": _export_excel(df),
                    "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                }
            )
        if "pdf" in attachment_options:
            attachments_payload.append(
                {
                    "filename": f"{export_basename}.pdf",
                    "content": _export_pdf(df, records, summary),
                    "content_type": "application/pdf",
                }
            )
        if "word" in attachment_options:
            attachments_payload.append(
                {
                    "filename": f"{export_basename}.docx",
                    "content": _export_word(df, records, summary),
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }
            )

        with st.spinner("Sending email..."):
            email_result = monitor.send_email(
                records,
                recipients=recipients,
                attachments=attachments_payload,
            )
        if email_result:
            st.success("Email sent.")
        elif not recipients:
            st.warning("Please add at least one recipient.")
        else:
            st.error("Email could not be sent. Check configuration and logs.")
